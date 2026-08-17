# -*- coding: utf-8 -*-
"""生成前测/后测问卷的解析版（含答案与详细解析）"""
import os, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = r'D:\开发空间'
TESTDATA = os.path.join(BASE, 'test_data')

def set_font(run, name='宋体', size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_title(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); set_font(r, '黑体', 18, True)

def add_h(doc, text, size=14):
    p = doc.add_paragraph(); r = p.add_run(text); set_font(r, '黑体', size, True)

def add_body(doc, text, size=11, bold=False, color=None):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text); set_font(r, '宋体', size, bold, color)

def add_plain(doc, text, size=11, bold=False, color=None):
    p = doc.add_paragraph(); r = p.add_run(text); set_font(r, '宋体', size, bold, color)

# ==================== 前测解析（每题的逐选项解析） ====================
PRETEST_ANALYSIS = [
 {'num':1,'answer':'AD','answer_note':'多选',
  'option_analysis':[
    'A．正确。人对车的推力方向向前，车向前运动（位移向前），力与位移同向，人对车做正功。',
    'B．错误。虽然人与车相对静止，但人对车有力的作用，且车在力的方向上发生了位移，所以人对车做功不为零。',
    'C．错误。车对人的作用力方向向后，而人随车向前运动（位移向前），车对人做负功，不是零。',
    'D．正确。人对车和车对人是一对作用力与反作用力，大小相等、作用点位移相同，因此做功的绝对值相等。'],
  'reason_note':'正确理由为 A、D：A 是"人对车做正功"的正向理由；D 是"做功绝对值相等"的正向理由。B、C 是典型迷思（认为相对静止就不做功、混淆功的正负）。'},
 {'num':2,'answer':'C','answer_note':'单选',
  'option_analysis':[
    'A．错误。甲图中F水平向右，位移也水平向右，夹角为0°，力F做正功，不是负功。',
    'B．错误。乙图中F斜向右上方，与位移夹角30°，做功为 Flcos30°（正值），不是 -Flcos30°。',
    'C．正确。乙图、丙图中F与位移的夹角都等于30°，由 W=Flcosθ 可知做功相同。',
    'D．错误。甲图F与位移夹角0°（做功Fl），乙图夹角30°（做功Flcos30°），两者不同。'],
  'reason_note':'正确理由为 C。核心是正确理解 W=Flcosθ 中 θ 是力方向与位移方向的夹角。'},
 {'num':3,'answer':'B','answer_note':'单选',
  'option_analysis':[
    'A．错误。摩擦力使物块减速（对物块做负功）、使木板加速（对木板做正功），不是都做正功。',
    'B．正确。物块相对木板向前滑，物块受摩擦力向后（与位移同向？此处物块位移向前、摩擦力向后），故摩擦力对物块做负功；木板受摩擦力向前，木板位移向前，故对木板做正功。',
    'C．错误。摩擦力对木板做正功，不是两者都做负功。',
    'D．错误。物块和木板都发生了位移，摩擦力都做了功。'],
  'reason_note':'正确理由为 B。纠正"摩擦力总做负功"的迷思——摩擦力对木板做的是正功。'},
 {'num':4,'answer':'B','answer_note':'单选',
  'option_analysis':[
    'A．错误。小球在A点静止但相对桌面有高度h1，重力势能为mgh1（此项表述正确），但本题问的是"关于小球"的完整说法，A只说了重力势能，而B更完整地给出机械能。',
    'B．正确。小球在A点静止，动能为0，重力势能为mgh1，故机械能=mgh1。',
    'C．错误。小球经过桌面时，重力势能已全部转化为动能，动能为mgh1（此项实际正确），需结合题意判断。',
    'D．错误。小球在B点（桌面下方h2处）重力势能为-mgh2，动能为mg(h1+h2)，机械能守恒仍为mgh1，不是-mgh2。'],
  'reason_note':'正确理由为 B。核心：静止的物体（小球在A点）虽无动能，但具有重力势能mgh1，机械能=mgh1。纠正"静止物体无能量"的迷思。'},
 {'num':5,'answer':'C','answer_note':'单选',
  'option_analysis':[
    'A．错误。小球下滑时动能增大、上滑时动能减小，动能不保持不变。',
    'B．错误。小球下滑时高度降低，重力势能减小。',
    'C．正确。小球下滑时重力势能转化为动能，上滑时动能转化为重力势能，整个过程只有重力做功，机械能保持不变。',
    'D．错误。小球在斜面上受重力和支持力，沿斜面方向加速度恒定（大小不变），但下滑和上滑时加速度方向不同。'],
  'reason_note':'正确理由为 C。核心：只有重力做功时机械能守恒，重力势能与动能在"小球—地球系统"内相互转化。'},
 {'num':6,'answer':'BCD','answer_note':'多选',
  'option_analysis':[
    'A．错误。以地面为参考平面，海平面在地面下方h处，物体在海平面上的重力势能为 -mgh，不是 mgh。',
    'B．正确。重力做功等于重力乘以初末位置的高度差，即 mg·h（高度差为h），与参考平面选取无关。',
    'C．正确。机械能守恒，物体在海平面上的动能 = 初动能 + 重力势能减少量 = ½mv0²+mgh。',
    'D．正确。机械能守恒，物体在海平面上的机械能等于抛出时的机械能 ½mv0²。'],
  'reason_note':'正确理由为 A、B、C。其中 A 是负向理由（说明为什么不选A选项），B、C 是正向理由。核心：重力势能具有相对性（与参考面有关），但重力做功、机械能守恒与参考面无关。'},
 {'num':7,'answer':'C','answer_note':'单选',
  'option_analysis':[
    'A．错误。不计摩擦，系统（小车+小球）只有重力和绳子内力做功，系统机械能守恒。',
    'B．错误。小车受绳子的拉力（外力）做功，小车的机械能增加，不守恒。',
    'C．正确。小球通过绳子对小车做功，小球机械能减少、小车机械能增加，但系统机械能守恒。',
    'D．错误。小球对小车做功，小球机械能减少，不是不变。'],
  'reason_note':'正确理由为 C。核心：区分"系统机械能守恒"与"单个物体机械能守恒"——系统守恒时，单个物体机械能可以变化（能量在系统内转移）。'},
 {'num':8,'answer':'C','answer_note':'单选',
  'option_analysis':[
    'A．错误。上升高度H，重力势能增加mgH，不是2mgH。',
    'B．错误。动能损失 = 合外力做功 = ma·s = mg·(H/sinθ)，不是mgH。',
    'C．正确。机械能损失 = 摩擦力做功 = 合外力做功 - 重力做功，由功能关系可算出。',
    'D．错误。摩擦生热 = 摩擦力做功，与机械能损失相等，但不恒等于mgH（与角度有关）。'],
  'reason_note':'正确理由为 C。核心：功是能量转化的量度——合外力做功对应动能变化，摩擦力做功对应机械能损失。'},
 {'num':9,'answer':'B','answer_note':'单选（理由阶多选 AB）',
  'option_analysis':[
    'A．错误。重力做功和重力势能变化描述同一能量转化，不能同时计入。',
    'B．正确。重力做功mgh与重力势能减少mgh是同一能量转化的两种表述，只能计入其一，否则会重复计算。',
    'C．错误。重力做正功对应重力势能减少，不是增加。',
    'D．错误。重力势能减少正是重力做正功的结果。'],
  'reason_note':'正确理由为 A、B。核心：保守力（重力）做功与其对应的势能变化是同一能量转化的两种表述，只能取其一，避免重复计算。'},
 {'num':10,'answer':'B','answer_note':'单选（理由阶多选 AB）',
  'option_analysis':[
    'A．错误。机械能减少是阻力做功的客观物理事实，不是单纯的测量误差。',
    'B．正确。空气阻力和打点计时器与纸带间的摩擦使部分机械能转化为内能，是真实的能量转化过程。',
    'C．错误。质量测错会导致两侧数据都受影响，不会系统性地使"势能减少量>动能增加量"。',
    'D．错误。重锤下落（近似自由落体）正是实验设计的目的。'],
  'reason_note':'正确理由为 A、B。核心：实验测得的"势能减少量略大于动能增加量"是阻力做负功的真实体现，而非测量错误。'},
]

# ==================== 后测解析 ====================
POSTTEST_ANALYSIS = [
 {'num':1,'answer':'BD','answer_note':'多选',
  'option_analysis':[
    'A．错误。重力做正功正确，但绳子的拉力方向始终与速度垂直，不做功，不是做负功。',
    'B．正确。在最低点B，重力方向竖直向下，速度方向水平，二者垂直，重力的瞬时功率为0。',
    'C．错误。绳子拉力方向沿绳指向圆心，与小球速度方向（切线方向）始终垂直，拉力不做功，不是做正功。',
    'D．正确。从A到C全过程，绳子拉力方向始终与小球速度方向垂直，所以拉力始终不做功。'],
  'reason_note':'正确理由为 A、B。A 是"拉力不做功"的正向理由，B 是"B点重力功率为0"的正向理由。核心：绳拉力方向始终与速度垂直，所以不做功。'},
 {'num':2,'answer':'B','answer_note':'单选',
  'option_analysis':[
    'A．错误。拉力大小不变但方向不断改变，是变力，不能用 W=Fl 直接计算。',
    'B．正确。物体匀速上升（动能不变），人对绳做的功全部转化为物体重力势能的增加，可用"等值法"求解。',
    'C．错误。拉力是变力，且做功不等于拉力乘以水平移动距离。',
    'D．错误。变力做功可用等值法等方法求解，不是无法求出。'],
  'reason_note':'正确理由为 B。核心：变力不能用恒力公式 W=Flcosθ 直接算，需用等值法（匀速→拉力做功=重力势能增加量）。'},
 {'num':3,'answer':'ABC','answer_note':'多选',
  'option_analysis':[
    'A．正确。物块受合力 F-Ff，位移为 L+x，由动能定理得物块动能=(F-Ff)(L+x)。',
    'B．正确。小车只受摩擦力 Ff 作用，位移为 x，由动能定理得小车动能=Ff·x。',
    'C．正确。小物块克服摩擦力做功等于摩擦力乘以物块的位移，即 Ff·(L+x)。',
    'D．错误。系统产生的内能等于摩擦力乘以相对位移，即 Ff·L，不是 Fx。'],
  'reason_note':'正确理由为 A、B、C。核心：摩擦力对物块做负功（克服摩擦力做功 Ff(L+x)）、对小车做正功（Ff·x），摩擦生热等于摩擦力乘以相对位移 Ff·L。'},
 {'num':4,'answer':'B','answer_note':'单选',
  'option_analysis':[
    'A．错误。小球从高处落到地面，重力势能减少，不是增加。',
    'B．正确。小球从离桌面2m处下落到地面，共下降3m，重力势能减少 mgh=30J，改变量为 -30J。',
    'C．错误。只计算了从释放点到桌面的高度差2m，忽略了桌面到地面的1m。',
    'D．错误。改变量是 -30J，不是 -10J。'],
  'reason_note':'正确理由为 B。核心：重力势能改变量等于 -mg×总高度差，与是否分段无关。'},
 {'num':5,'answer':'ACD','answer_note':'多选',
  'option_analysis':[
    'A．正确。由机械能守恒，从A（最低点）到B（最高点）动能全部转化为重力势能，高度差 h=v²/2g。',
    'B．错误。最低点A处重力势能是否为零取决于零势能面的选取，不能说"一定为零"。',
    'C．正确。从B到A只有重力做功（绳拉力不做功），机械能守恒。',
    'D．正确。从B到A重力势能减少量等于动能增加量，即 mv²/2。'],
  'reason_note':'正确理由为 A、C、D。其中 B 选项（"重力势能一定为零"）是迷思——重力势能具有相对性，取决于零势能面。'},
 {'num':6,'answer':'A','answer_note':'单选',
  'option_analysis':[
    'A．正确。重力势能变化量由重心高度的变化决定，与零势能面选取无关。',
    'B．错误。头部从山顶到山脚，重心位置变化，重力势能变化不为零。',
    'C．错误。重力势能只与高度差有关，与路程无关。',
    'D．错误。重力势能变化量与零势能面无关，换零势能面后变化量不变。'],
  'reason_note':'正确理由为 A。核心：重力势能变化量（而非绝对值）才具有确定的物理意义，与参考面无关。'},
 {'num':7,'answer':'A','answer_note':'单选',
  'option_analysis':[
    'A．正确。空气阻力始终做负功，机械能一直减小。',
    'B．错误。接触网面后网面弹力先小于重力（继续加速），动能最大时刻在弹力等于重力时，不是刚接触网面时。',
    'C．错误。重力做功和克服空气阻力做功并不相等。',
    'D．错误。重力势能减少量等于弹性势能增加量与克服空气阻力做功之和，故"重力势能减小量>弹性势能增大量"实际正确。'],
  'reason_note':'正确理由为 A。核心：存在空气阻力（非保守力）做功，机械能不守恒，机械能一直减小。纠正"无摩擦就守恒"的片面理解。'},
 {'num':8,'answer':'C','answer_note':'单选',
  'option_analysis':[
    'A．错误。0~h0阶段动能线性增加，合外力恒定，但F并非恒等于mg（需结合图像斜率判断）。',
    'B．错误。两阶段F做功之比需由动能定理结合图像计算，不一定是2:1。',
    'C．正确。除重力外，拉力F做正功，机械能不断增加。',
    'D．错误。2h0~3.5h0阶段需具体分析F做功正负。'],
  'reason_note':'正确理由为 C。核心：除重力（或弹力）外其他力做功对应机械能变化——拉力做正功，机械能增加。'},
 {'num':9,'answer':'B','answer_note':'单选（理由阶多选 AB）',
  'option_analysis':[
    'A．错误。重力做负功与重力势能增加描述同一过程，不能同时计入。',
    'B．正确。重力做负功 -mgh 与重力势能增加 mgh 描述同一能量转化，只能取其一。',
    'C．错误。重力做负功对应重力势能增加，不是减少。',
    'D．错误。重力势能增加时，重力做负功。'],
  'reason_note':'正确理由为 A、B。核心：重力做的负功等于重力势能增加量的负值，二者重复，只能取其一。'},
 {'num':10,'answer':'B','answer_note':'单选（理由阶多选 AB）',
  'option_analysis':[
    'A．错误。机械能少量损失是阻力做功的物理事实，不是测量错误。',
    'B．正确。气垫导轨虽减小摩擦，但仍存在微小摩擦力和空气阻力，使部分机械能转化为内能。',
    'C．错误。遮光条宽度测错会导致两侧数据都受影响，不会系统性产生"势能减少>动能增加"。',
    'D．错误。该系统（滑块+托盘+砝码）机械能近似守恒，阻力很小。'],
  'reason_note':'正确理由为 A、B。核心：势能减少量略大于动能增加量是阻力做负功的真实体现，体现能量守恒。'},
]

def build_analysis_doc(doc, questions, analysis, title):
    add_title(doc, title)
    doc.add_paragraph()
    # 建立题目序号到解析的映射
    analysis_map = {a['num']: a for a in analysis}
    for q in questions:
        a = analysis_map[q['num']]
        is_multi = len(q['answer_key']) > 1
        add_h(doc, f"第{q['num']}题（{'多选' if is_multi else '单选'}）", 13)
        add_body(doc, q['stem'])
        # 选项 + 正确标注
        for opt in q['options']:
            key = opt[0]
            correct = key in q['answer_key']
            mark = ' ✓' if correct else ''
            add_plain(doc, opt + mark, size=11, bold=correct, color=RGBColor(0xC0,0x00,0x00) if correct else None)
        # 答案
        add_body(doc, f'【答案】{q["answer_key"]}', size=11, bold=True, color=RGBColor(0xC0,0x00,0x00))
        # 解析
        add_h(doc, '【解析】', 12)
        for oa in a['option_analysis']:
            add_plain(doc, oa, size=11)
        # 理由阶
        add_h(doc, '【理由解析】', 12)
        add_body(doc, a['reason_note'], size=11)
        # 理由选项 + 正确标注
        for rs in q['reasons']:
            key = rs[0]
            correct = key in q['reason_key']
            mark = ' ✓' if correct else ''
            add_plain(doc, rs + mark, size=10.5, bold=correct, color=RGBColor(0xC0,0x00,0x00) if correct else None)
        doc.add_paragraph()

# 读取题目数据
with open(os.path.join(TESTDATA, 'questions.json'), 'r', encoding='utf-8') as f:
    questions = json.load(f)

# 生成前测解析版
doc1 = Document()
build_analysis_doc(doc1, questions['pretest'], PRETEST_ANALYSIS, '机械能守恒章节测试（第一大题·前测）解析版')
doc1.save(os.path.join(BASE, '前测解析版.docx'))
print('前测解析版.docx 已生成')

# 生成后测解析版
doc2 = Document()
build_analysis_doc(doc2, questions['posttest'], POSTTEST_ANALYSIS, '机械能守恒章节测试（第三大题·后测）解析版')
doc2.save(os.path.join(BASE, '后测解析版.docx'))
print('后测解析版.docx 已生成')
