# -*- coding: utf-8 -*-
"""四阶问卷批改与Word报告生成的公共工具（供 Agent2~5 使用）"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def set_font(run, name='宋体', size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def judge_tier_wrong(answer, reason, answer_key, reason_key):
    """判断四阶题是否答错：答案阶或理由阶任一出错即错（支持单选/多选）"""
    answer_wrong = (sorted(answer or '') != sorted(answer_key or ''))
    reason_wrong = (sorted(reason or '') != sorted(reason_key or ''))
    return {
        'answer_wrong': answer_wrong,
        'reason_wrong': reason_wrong,
        'is_wrong': answer_wrong or reason_wrong,
    }


def grade_four_tier(questions, answers, key_prefix):
    """
    批改一组四阶题，返回逐题结果列表。

    questions: 题目列表（含 num/miscon/stem/answer_key/reason_key）
    answers: 答题数据 dict，key 为 key_prefix + (num-1)
    key_prefix: 'q'（前测）或 'p'（后测）
    """
    results = []
    for q in questions:
        student = answers.get(f'{key_prefix}{q["num"]-1}') or {}
        a = student.get('answer', '')
        r = student.get('reason', '')
        conf1 = student.get('conf1', '')
        conf2 = student.get('conf2', '')
        j = judge_tier_wrong(a, r, q['answer_key'], q['reason_key'])
        results.append({
            'num': q['num'],
            'miscon': q['miscon'],
            'stem': q['stem'],
            'answer': a, 'reason': r,
            'conf1': conf1, 'conf2': conf2,
            'answer_key': q['answer_key'],
            'reason_key': q['reason_key'],
            'is_wrong': j['is_wrong'],
            'answer_wrong': j['answer_wrong'],
            'reason_wrong': j['reason_wrong'],
        })
    return results


def grade_correction(questions, answers):
    """批改纠错题（单选，key 为 c1~cN）"""
    results = []
    for i, q in enumerate(questions):
        student = answers.get(f'c{i+1}') or {}
        a = student.get('answer', '')
        correct = (a == q['answer'])
        results.append({
            'num': i + 1,
            'miscon': q.get('miscon', ''),
            'name': q.get('name', ''),
            'stem': q['stem'],
            'answer': a,
            'answer_key': q['answer'],
            'is_wrong': not correct,
        })
    return results


def _build_report(title, student_name, student_id, sections, output_path):
    """
    生成 Word 报告。sections 为列表，每项为 {'heading': str, 'table': {'headers':[...], 'rows':[[...],...]}, 'texts': [str, ...]}
    """
    doc = Document()
    # 页面设置
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_font(r, '黑体', 18, True)

    # 学生信息
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'学生：{student_name}　学号：{student_id}')
    set_font(r, '宋体', 12)

    for s in sections:
        # 小标题
        p = doc.add_paragraph()
        r = p.add_run(s['heading'])
        set_font(r, '黑体', 14, True)

        # 表格
        if 'table' in s and s['table']:
            t = s['table']
            table = doc.add_table(rows=1, cols=len(t['headers']))
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            for i, h in enumerate(t['headers']):
                hdr[i].text = str(h)
                for para in hdr[i].paragraphs:
                    for run in para.runs:
                        set_font(run, '宋体', 10, True)
            for row in t['rows']:
                cells = table.add_row().cells
                for i, cell_val in enumerate(row):
                    cells[i].text = str(cell_val)
                    for para in cells[i].paragraphs:
                        for run in para.runs:
                            set_font(run, '宋体', 10)

        # 文本段落
        if 'texts' in s:
            for text in s['texts']:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0.74)
                r = p.add_run(text)
                set_font(r, '宋体', 11)

        doc.add_paragraph()

    doc.save(output_path)
    return output_path


def collect_misconceptions(results):
    """从批改结果中提取错题对应的迷思概念（去重）"""
    miscon_set = set()
    for r in results:
        if r['is_wrong']:
            for m in r['miscon'].split('+'):
                miscon_set.add(m)
    return sorted(miscon_set)
