# -*- coding: utf-8 -*-
"""
统一分析模块 —— 批改答题、生成Word诊断报告
供 server.py（实时分析）和独立脚本（离线分析）共用
"""
import json, os, re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# 项目根目录
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
KEYS_PATH = os.path.join(BASE_DIR, 'test_data', 'answer_keys.json')
REPORT_DIR = os.path.join(BASE_DIR, 'reports')
os.makedirs(REPORT_DIR, exist_ok=True)


def load_answer_keys():
    """加载答案键"""
    with open(KEYS_PATH, 'r', encoding='utf-8') as f:
        return json.load(f)


def find_answer_data(section):
    """查找最新的答题数据文件"""
    answer_dir = os.path.join(BASE_DIR, 'answer_data', section)
    files = []
    # 在 answer_data/xxx/ 中查找
    if os.path.exists(answer_dir):
        for f in os.listdir(answer_dir):
            if f.startswith('answer_') and f.endswith('.json'):
                files.append((os.path.getmtime(os.path.join(answer_dir, f)), os.path.join(answer_dir, f)))
    # 也在项目根目录查找
    for f in os.listdir(BASE_DIR):
        if f.startswith('answer_' + section) and f.endswith('.json'):
            fp = os.path.join(BASE_DIR, f)
            files.append((os.path.getmtime(fp), fp))
    if files:
        files.sort(reverse=True)
        return files[0][1]
    return None


def grade_single(key_answer, student_answer):
    """批改单选题"""
    sa = student_answer.get('v', '') if isinstance(student_answer, dict) else str(student_answer)
    correct = (str(sa).strip().upper() == str(key_answer).strip().upper())
    return {
        'correct': correct,
        'student': str(sa),
        'expected': str(key_answer)
    }


def grade_multi(key_answer, student_answer):
    """批改多选题"""
    # 提取学生答案 - 支持多种格式
    if isinstance(student_answer, dict):
        sa = student_answer.get('v', student_answer.get('value', []))
    elif isinstance(student_answer, list):
        sa = student_answer
    else:
        sa = []
    
    # 确保是列表
    if not isinstance(sa, list):
        sa = [sa] if sa else []
    
    key_set = set(key_answer)
    stu_set = set(sa)
    
    if stu_set == key_set:
        status, ratio = 'correct', 1.0
    elif len(stu_set) == 0:
        status, ratio = 'empty', 0.0
    elif stu_set.issubset(key_set):
        status, ratio = 'partial', 0.5
    elif key_set.issubset(stu_set):
        status, ratio = 'over', 0.5
    else:
        status, ratio = 'wrong', 0.0
    
    return {
        'correct': status == 'correct',
        'status': status,
        'ratio': ratio,
        'student': sorted(stu_set),
        'expected': sorted(key_set)
    }


def grade_text(key_subs, student_value):
    """批改计算/填空类题目（基于关键词匹配）"""
    # 提取学生答案的实际文本值
    if isinstance(student_value, dict):
        sv = student_value.get('v', student_value)
    else:
        sv = student_value
    
    # sv可能是dict(多子题)或str(单文本)
    if isinstance(sv, str):
        sv = {'0': sv}
    
    sub_results = []
    total_ratio = 0
    
    for i, sub_key in enumerate(key_subs):
        student_text = str(sv.get(str(i), '')).strip()
        keywords = sub_key.get('keywords', [])
        
        if not student_text:
            sub_results.append({'ratio': 0, 'status': 'empty', 'student': '（未作答）'})
            continue
        
        # 关键词匹配评分
        matched = 0
        for kw in keywords:
            if kw.lower() in student_text.lower():
                matched += 1
        
        if len(keywords) > 0:
            ratio = matched / len(keywords)
        else:
            ratio = 0.5 if len(student_text) > 5 else 0.3  # 有内容就给基础分
        
        if ratio >= 0.8:
            status = 'good'
        elif ratio >= 0.4:
            status = 'partial'
        else:
            status = 'poor'
        
        sub_results.append({
            'ratio': ratio,
            'status': status,
            'student': student_text[:80] + ('...' if len(student_text) > 80 else ''),
            'keywords_matched': matched,
            'keywords_total': len(keywords)
        })
        total_ratio += ratio
    
    avg_ratio = total_ratio / max(len(key_subs), 1)
    return {'sub_results': sub_results, 'ratio': avg_ratio}


def analyze_section(section_key, answer_data=None):
    """
    分析某一大题的答题情况
    section_key: 'pretest' | 'correction' | 'posttest'
    answer_data: 作答数据字典，如果为None则自动查找最新文件
    
    返回: {
        'results': [...],  # 每题批改结果
        'total_score': int,
        'max_score': int,
        'summary': {...}
    }
    """
    keys = load_answer_keys()
    section_keys = keys.get(section_key, {})
    sec_name = section_keys.get('name', section_key)
    question_keys = section_keys.get('questions', {})
    
    # 加载答题数据
    if answer_data is None:
        filepath = find_answer_data(section_key)
        if not filepath:
            return None
        with open(filepath, 'r', encoding='utf-8') as f:
            raw = json.load(f)
        student_answers = raw.get('answers', {})
    else:
        student_answers = answer_data
    
    # 按题目顺序分析
    results = []
    total_score = 0
    max_score = 0
    
    # 对答案键中的每道题进行批改
    sorted_keys = sorted(question_keys.keys(), key=lambda x: int(x[1:]) if x[1:].isdigit() else 999)
    
    # 建立学生答案的索引
    student_map = {}  # 按题号(1-based)映射到学生答案
    for k, v in student_answers.items():
        k_clean = k.lstrip('q')
        if '_q' in k:
            parts = k.split('_q')
            if len(parts) == 2:
                sn = parts[0].lstrip('s')
                qn = parts[1]
                if sn == section_key[-1] or sn == str({'pretest':'1','correction':'2','posttest':'3'}.get(section_key, '0')):
                    student_map[str(int(qn))] = v
        elif k_clean.isdigit():
            # QB全局索引(0-based) → 转换为1-based题号
            student_map[str(int(k_clean) + 1)] = v
    
    # 兜底：如果还匹配不上，直接用原始key
    if not student_map:
        for k, v in student_answers.items():
            clean = k.lstrip('q')
            if clean.isdigit():
                student_map[str(int(clean) + 1)] = v
    
    for idx, qkey in enumerate(sorted_keys):
        key = question_keys[qkey]
        qtype = key.get('type', '')
        qscore = key.get('score', 6)
        max_score += qscore
        
        # 查找学生答案（尝试多种key匹配）
        # 先尝试直接匹配 answer key 的 qkey
        student_ans = student_answers.get(qkey)
        # 如果失败，尝试用题号从 student_map 中查找
        if student_ans is None and student_map:
            qnum = str(idx + 1)  # 1-based题目序号
            student_ans = student_map.get(qnum)
        
        result = {
            'qkey': qkey,
            'qnum': idx + 1,
            'type': qtype,
            'max_score': qscore,
            'score': 0,
            'correct': False,
            'student_answer': None,
            'expected_answer': None,
            'detail': ''
        }
        
        if not student_ans:
            result['detail'] = '未作答'
            results.append(result)
            continue
        
        if qtype == '单选':
            ka = key.get('answer', '')
            gr = grade_single(ka, student_ans)
            result['correct'] = gr['correct']
            result['score'] = qscore if gr['correct'] else 0
            result['student_answer'] = gr['student']
            result['expected_answer'] = gr['expected']
            result['detail'] = '正确' if gr['correct'] else f'错误（答案：{gr["student"]}，正确答案：{gr["expected"]}）'
            result['explanation'] = key.get('explanation', '')
        
        elif qtype == '多选':
            ka = key.get('answer', [])
            try:
                gr = grade_multi(ka, student_ans)
            except Exception as e:
                import sys
                sys.stderr.write(f'[GRADE ERROR] qkey={qkey} type=multi ans={student_ans} err={e}\n')
                gr = {'correct': False, 'status': 'error', 'ratio': 0, 'student': str(student_ans)[:50], 'expected': ka}
            result['correct'] = gr['correct']
            result['score'] = int(qscore * gr['ratio'])
            result['student_answer'] = gr['student']
            result['expected_answer'] = gr['expected']
            status_map = {'correct': '全对 ✓', 'partial': '漏选（得一半分）', 'over': '多选（得一半分）', 'wrong': '错误', 'empty': '未作答'}
            result['detail'] = status_map.get(gr['status'], '')
            result['explanation'] = key.get('explanation', '')
        
        elif qtype == '填空':
            ka = key.get('answer', '')
            gr = grade_single(ka, student_ans)
            result['correct'] = gr['correct']
            result['score'] = qscore if gr['correct'] else 0
            result['student_answer'] = gr['student']
            result['expected_answer'] = gr['expected']
            result['detail'] = '正确' if gr['correct'] else f'错误（答案：{gr["student"]}，正确答案：{gr["expected"]}）'
            result['explanation'] = key.get('explanation', '')
        
        elif qtype == '计算':
            sub_keys = key.get('sub_answers', [])
            gr = grade_text(sub_keys, student_ans)
            result['score'] = int(qscore * gr['ratio'])
            result['correct'] = gr['ratio'] >= 0.6
            result['sub_results'] = gr['sub_results']
            result['student_answer'] = [s['student'] for s in gr['sub_results']]
            result['detail'] = f'得分率 {gr["ratio"]*100:.0f}%'
            result['explanation'] = key.get('explanation', '')
        
        total_score += result['score']
        results.append(result)
    
    # 统计
    correct_count = sum(1 for r in results if r['correct'])
    ratio = total_score / max(max_score, 1) * 100
    
    if ratio >= 90: level = '优秀'
    elif ratio >= 75: level = '良好'
    elif ratio >= 60: level = '一般'
    else: level = '需加强'
    
    return {
        'section': section_key,
        'section_name': sec_name,
        'results': results,
        'total_score': total_score,
        'max_score': max_score,
        'ratio': ratio,
        'level': level,
        'correct_count': correct_count,
        'total_count': len(results),
        'timestamp': datetime.now()
    }


def generate_docx_report(analysis, output_name=None, output_dir=None):
    """
    生成Word诊断报告
    
    Args:
        analysis: analyze_section() 的返回结果
        output_name: 自定义文件名（不含扩展名）
        output_dir: 自定义输出目录（默认 REPORT_DIR）
    """
    if analysis is None:
        return None
    
    doc = Document()
    
    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # --- 标题 ---
    title = doc.add_heading('物理习题课·诊断分析报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run(f'测试模块：电磁感应章节习题课').font.size = Pt(12)
    
    # 学生信息（如果有）
    student_name = analysis.get('student_name', '')
    student_id = analysis.get('student_id', '')
    if student_name:
        info_para.add_run(f'\n学生：{student_name}  |  学号：{student_id}').font.size = Pt(11)
    info_para.add_run(f'\n{analysis["section_name"]}').font.size = Pt(12)
    
    # 基本信息表
    doc.add_paragraph()
    info_table = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ('生成时间', analysis['timestamp'].strftime('%Y年%m月%d日 %H:%M:%S')),
        ('分析内容', analysis['section_name']),
        ('题目数量', f'{analysis["total_count"]} 题'),
        ('总分', f'{analysis["total_score"]} / {analysis["max_score"]}（{analysis["ratio"]:.1f}%）'),
        ('评价等级', analysis['level'])
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = str(value)
        for cell in [info_table.cell(i, 0), info_table.cell(i, 1)]:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # --- 逐题诊断 ---
    doc.add_heading('一、逐题诊断', level=1)
    
    for i, r in enumerate(analysis['results']):
        # 题目标题
        status_icon = '✅' if r['correct'] else ('⚠️' if r['score'] > 0 else '❌')
        q_title = doc.add_heading(f'第{r["qnum"]}题 {status_icon}  [{r["type"]}] —— {r["score"]}/{r["max_score"]}分', level=2)
        
        # 学生答案 vs 标准答案
        detail_para = doc.add_paragraph()
        if r['type'] in ('单选', '多选', '填空'):
            detail_para.add_run('学生答案：').bold = True
            detail_para.add_run(str(r.get('student_answer', '未作答')))
            detail_para.add_run('\n标准答案：').bold = True
            detail_para.add_run(str(r.get('expected_answer', '')))
        elif r['type'] == '计算' and 'sub_results' in r:
            for si, sr in enumerate(r['sub_results']):
                sub_status = {'good': '✓', 'partial': '△', 'poor': '✗', 'empty': '○'}.get(sr['status'], '?')
                detail_para.add_run(f'\n({si+1}) {sub_status} {sr["student"]}')
        
        # 诊断说明
        diag_para = doc.add_paragraph()
        diag_para.add_run('诊断：').bold = True
        diag_para.add_run(r.get('detail', ''))
        
        # 解析
        if r.get('explanation'):
            exp_para = doc.add_paragraph()
            exp_para.add_run('解析：').bold = True
            exp_para.add_run(r['explanation'])
            exp_para.runs[-1].font.size = Pt(10)
            exp_para.runs[-1].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    # --- 知识模块分析 ---
    doc.add_heading('二、知识模块分析', level=1)
    
    correct_qs = [r for r in analysis['results'] if r['correct']]
    wrong_qs = [r for r in analysis['results'] if not r['correct'] and r['score'] == 0]
    partial_qs = [r for r in analysis['results'] if not r['correct'] and r['score'] > 0]
    
    p = doc.add_paragraph()
    p.add_run(f'完全正确：{len(correct_qs)} 题').bold = True
    p.add_run(f'\n部分正确：{len(partial_qs)} 题')
    p.add_run(f'\n完全错误：{len(wrong_qs)} 题')
    
    if correct_qs:
        p.add_run('\n\n✅ 已掌握的知识点：').bold = True
        for r in correct_qs:
            p.add_run(f'\n  · 第{r["qnum"]}题（{r["type"]}）')
    
    if partial_qs:
        p.add_run('\n\n⚠️ 部分掌握的知识点：').bold = True
        for r in partial_qs:
            p.add_run(f'\n  · 第{r["qnum"]}题（{r["type"]}）—— 得分 {r["score"]}/{r["max_score"]}')
    
    if wrong_qs:
        p.add_run('\n\n❌ 需要加强的知识点：').bold = True
        for r in wrong_qs:
            p.add_run(f'\n  · 第{r["qnum"]}题（{r["type"]}）')
    
    # --- 学习建议 ---
    doc.add_heading('三、学习建议', level=1)
    
    ratio = analysis['ratio']
    if ratio >= 90:
        advice = '学生对本部分知识掌握扎实，建议在后测中重点验证知识迁移能力。'
    elif ratio >= 75:
        advice = '学生知识掌握情况良好，个别知识点存在理解偏差，建议针对性复习后进行后测。'
    elif ratio >= 60:
        advice = '学生知识掌握一般，存在若干薄弱环节。建议重点复习以下内容：'
        for r in wrong_qs:
            advice += f'\n  · 第{r["qnum"]}题涉及的知识点'
    else:
        advice = '学生基础较为薄弱，建议系统复习本部分基础知识后再进行后测。重点关注：'
        for r in wrong_qs:
            advice += f'\n  · 第{r["qnum"]}题涉及的知识点'
    
    doc.add_paragraph(advice)
    
    # --- 页脚 ---
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run('—— 本报告由 WorkBuddy AI 诊断系统自动生成 ——').font.size = Pt(9)
    footer_para.runs[-1].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    # 保存
    if output_name is None:
        ts = analysis['timestamp'].strftime('%Y%m%d_%H%M%S')
        section_key = analysis['section']
        output_name = f'{section_key}_diagnosis_{ts}'
    
    save_dir = output_dir if output_dir else REPORT_DIR
    os.makedirs(save_dir, exist_ok=True)
    filepath = os.path.join(save_dir, f'{output_name}.docx')
    doc.save(filepath)
    return filepath


# ===== 命令行入口（兼容旧的独立分析脚本） =====
if __name__ == '__main__':
    import sys
    if len(sys.argv) < 2:
        print('用法: python analyzer.py <section>  (section: pretest/correction/posttest)')
        sys.exit(1)
    
    section = sys.argv[1]
    print(f'分析 {section}...')
    result = analyze_section(section)
    if result is None:
        print('未找到答题数据！')
        sys.exit(1)
    
    filepath = generate_docx_report(result)
    if filepath:
        print(f'报告已生成: {filepath}')
        print(f'总分: {result["total_score"]}/{result["max_score"]} ({result["ratio"]:.1f}%)')
        print(f'等级: {result["level"]}')
